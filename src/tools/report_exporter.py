from __future__ import annotations

from pathlib import Path
from typing import Any


class ReportExporter:
    """Export Markdown reports to docx and optionally pdf."""

    def __init__(self, reports_dir: Path) -> None:
        self.reports_dir = reports_dir
        self.reports_dir.mkdir(parents=True, exist_ok=True)

    def markdown_to_docx(self, markdown_path: Path, docx_path: Path | None = None) -> dict[str, Any]:
        docx_path = docx_path or markdown_path.with_suffix(".docx")
        try:
            from docx import Document
        except ImportError:
            return {
                "success": False,
                "path": str(docx_path),
                "warning": "DOCX export skipped because python-docx is unavailable.",
            }

        document = Document()
        for line in markdown_path.read_text(encoding="utf-8").splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("# "):
                document.add_heading(stripped[2:].strip(), level=0)
            elif stripped.startswith("## "):
                document.add_heading(stripped[3:].strip(), level=1)
            elif stripped.startswith("### "):
                document.add_heading(stripped[4:].strip(), level=2)
            elif stripped.startswith("|") and stripped.endswith("|"):
                document.add_paragraph(stripped)
            elif stripped.startswith("![") and "](" in stripped and stripped.endswith(")"):
                document.add_paragraph(stripped)
            elif stripped.startswith("```"):
                document.add_paragraph(stripped)
            else:
                document.add_paragraph(stripped)
        document.save(docx_path)
        return {"success": True, "path": str(docx_path)}

    def markdown_to_pdf(self, markdown_path: Path, pdf_path: Path | None = None) -> dict[str, Any]:
        pdf_path = pdf_path or markdown_path.with_suffix(".pdf")
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
        except ImportError:
            return {
                "success": False,
                "path": str(pdf_path),
                "warning": "PDF export skipped because dependency is unavailable.",
            }

        text = markdown_path.read_text(encoding="utf-8")
        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        width, height = A4
        y = height - 48
        c.setFont("Helvetica", 10)
        for line in text.splitlines():
            if y < 48:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = height - 48
            c.drawString(48, y, line[:110])
            y -= 14
        c.save()
        return {"success": True, "path": str(pdf_path)}

    def export_all(
        self,
        markdown_path: Path,
        export_docx: bool = True,
        export_pdf: bool = False,
    ) -> dict[str, Any]:
        result = {"markdown": str(markdown_path), "docx": None, "pdf": None, "warnings": []}
        if export_docx:
            docx = self.markdown_to_docx(markdown_path)
            result["docx"] = docx
            if docx.get("warning"):
                result["warnings"].append(docx["warning"])
        if export_pdf:
            pdf = self.markdown_to_pdf(markdown_path)
            result["pdf"] = pdf
            if pdf.get("warning"):
                result["warnings"].append(pdf["warning"])
        return result
